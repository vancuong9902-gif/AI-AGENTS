from __future__ import annotations

import io
import random
import zipfile
from dataclasses import dataclass
from itertools import combinations
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from sqlalchemy.orm import Session

from app.models.document_topic import DocumentTopic
from app.models.question import Question


@dataclass
class GeneratedVariant:
    code: str
    questions: list[dict[str, Any]]


def _set_two_columns(document: Document) -> None:
    section = document.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    col = cols[0] if cols else None
    if col is None:
        col = sectPr.makeelement(qn("w:cols"))
        sectPr.append(col)
    col.set(qn("w:num"), "2")


def _set_doc_font(run) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def _topic_names(db: Session, topic_ids: list[int]) -> list[str]:
    rows = (
        db.query(DocumentTopic.teacher_edited_title, DocumentTopic.title)
        .filter(DocumentTopic.id.in_([int(x) for x in (topic_ids or [])]))
        .all()
    )
    return [str(a or b or "").strip() for a, b in rows if str(a or b or "").strip()]


def _question_bank(db: Session, topic_names: list[str], limit: int = 400) -> list[Question]:
    q = db.query(Question)
    if topic_names:
        from sqlalchemy import or_

        q = q.filter(or_(*[Question.stem.ilike(f"%{t}%") for t in topic_names]))
    return q.limit(int(limit)).all()


def _variant_overlap_ratio(a: list[int], b: list[int]) -> float:
    sa, sb = set(a), set(b)
    if not sa:
        return 0.0
    return len(sa & sb) / len(sa)


def _pick_variants(pool: list[Question], *, num_variants: int, questions_per_exam: int, min_unique_ratio: float) -> list[GeneratedVariant]:
    if len(pool) < questions_per_exam:
        raise ValueError("Not enough questions to generate exam variants")

    rng = random.Random(42)
    variants: list[GeneratedVariant] = []
    for idx in range(num_variants):
        for _ in range(200):
            picked = rng.sample(pool, questions_per_exam)
            picked_ids = [int(q.id) for q in picked]
            ok = True
            for prev in variants:
                overlap = _variant_overlap_ratio([int(x["id"]) for x in prev.questions], picked_ids)
                unique_ratio = 1.0 - overlap
                if unique_ratio < float(min_unique_ratio):
                    ok = False
                    break
            if ok:
                q_payload = []
                for q in picked:
                    options = list(q.options or [])
                    correct_idx = int(getattr(q, "correct_index", 0) or 0)
                    indexed = list(enumerate(options))
                    rng.shuffle(indexed)
                    shuffled = [opt for _, opt in indexed]
                    new_correct = 0
                    for ni, (oi, _) in enumerate(indexed):
                        if oi == correct_idx:
                            new_correct = ni
                            break
                    q_payload.append({"id": int(q.id), "stem": str(q.stem), "options": shuffled, "correct_index": new_correct})
                variants.append(GeneratedVariant(code=f"{idx+1:03d}", questions=q_payload))
                break
        else:
            raise ValueError("Could not satisfy min_unique_ratio across variants")
    return variants


def _build_exam_docx(*, variant: GeneratedVariant, exam_title: str, school_name: str, subject: str) -> bytes:
    doc = Document()
    _set_two_columns(doc)

    p = doc.add_paragraph()
    r = p.add_run(f"{school_name}\nMã đề: {variant.code} | Môn: {subject} | Thời gian: 45 phút")
    _set_doc_font(r)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(exam_title)
    _set_doc_font(r2)
    r2.bold = True

    p3 = doc.add_paragraph()
    r3 = p3.add_run("Họ tên: ....................... | Lớp: ........ | Ngày: ........")
    _set_doc_font(r3)

    for i, q in enumerate(variant.questions, start=1):
        pq = doc.add_paragraph()
        rq = pq.add_run(f"Câu {i}: {q['stem']}")
        _set_doc_font(rq)

        opts = q.get("options") or []
        labels = ["A", "B", "C", "D"]
        for oi, opt in enumerate(opts[:4]):
            po = doc.add_paragraph()
            ro = po.add_run(f"{labels[oi]}. {opt}")
            _set_doc_font(ro)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _build_answer_key_docx(variants: list[GeneratedVariant], exam_title: str) -> bytes:
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run(f"ĐÁP ÁN - {exam_title}")
    _set_doc_font(r)
    r.bold = True
    labels = ["A", "B", "C", "D"]
    for v in variants:
        pv = doc.add_paragraph()
        rv = pv.add_run(f"Mã đề {v.code}")
        _set_doc_font(rv)
        rv.bold = True
        for i, q in enumerate(v.questions, start=1):
            pa = doc.add_paragraph()
            ra = pa.add_run(f"Câu {i}: {labels[int(q.get('correct_index', 0))]}")
            _set_doc_font(ra)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_exam_docx_zip(
    db: Session,
    *,
    classroom_id: int,
    topic_ids: list[int],
    num_variants: int,
    questions_per_exam: int,
    exam_type: str,
    difficulty_distribution: dict[str, int] | None,
    include_answer_key: bool,
    exam_title: str,
    school_name: str,
    subject: str,
    min_unique_ratio: float = 0.6,
) -> bytes:
    _ = (classroom_id, exam_type, difficulty_distribution)
    topic_names = _topic_names(db, topic_ids)
    pool = _question_bank(db, topic_names)
    variants = _pick_variants(
        pool,
        num_variants=int(num_variants),
        questions_per_exam=int(questions_per_exam),
        min_unique_ratio=float(min_unique_ratio),
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for v in variants:
            zf.writestr(f"de_{v.code}.docx", _build_exam_docx(variant=v, exam_title=exam_title, school_name=school_name, subject=subject))
        if include_answer_key:
            zf.writestr("dap_an.docx", _build_answer_key_docx(variants, exam_title))
    buf.seek(0)
    return buf.getvalue()
