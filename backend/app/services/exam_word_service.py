from __future__ import annotations

import copy
import io
import random
import zipfile
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm


def generate_exam_word(questions: list[dict[str, Any]], num_versions: int, include_answer_key: bool) -> tuple[bytes, str]:
    all_docs: list[bytes] = []

    for version_num in range(1, int(num_versions) + 1):
        doc = Document()

        section = doc.sections[0]
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)

        header = doc.add_heading(f"BÀI KIỂM TRA - ĐỀ SỐ {version_num}", 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("Họ và tên: .............................................")
        doc.add_paragraph("Thời gian: 45 phút  |  Ngày: ___/___/______")
        doc.add_paragraph("─" * 60)

        version_questions = copy.deepcopy(questions)
        random.shuffle(version_questions)

        for i, q in enumerate(version_questions, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"Câu {i}: {q['question_text']}")
            run.bold = True

            if q["type"] == "multiple_choice":
                options = list(zip(["A", "B", "C", "D"], q.get("options") or []))
                random.shuffle(options)
                for letter, opt in options:
                    doc.add_paragraph(f"   {letter}. {opt}")
            else:
                for _ in range(4):
                    doc.add_paragraph("   _______________________________________________")

            doc.add_paragraph("")

        if include_answer_key:
            doc.add_page_break()
            doc.add_heading(f"ĐÁP ÁN - ĐỀ SỐ {version_num}", 1)
            for i, q in enumerate(version_questions, 1):
                doc.add_paragraph(f"Câu {i}: {q['correct_answer']}")

        buf = io.BytesIO()
        doc.save(buf)
        all_docs.append(buf.getvalue())

    if int(num_versions) == 1:
        return all_docs[0], "exam.docx"

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w") as zf:
        for i, doc_bytes in enumerate(all_docs, 1):
            zf.writestr(f"de_thi_so_{i}.docx", doc_bytes)
    return zip_buf.getvalue(), "de_thi.zip"
