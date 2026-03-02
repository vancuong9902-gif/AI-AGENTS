from __future__ import annotations

import json
import os
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.oxml.ns import qn

from app.services.vietnamese_font_fix import fix_vietnamese_text, VIETNAMESE_DOCX_FONT


def _set_times_new_roman(doc: Document) -> None:
    style = doc.styles['Normal']
    style.font.name = VIETNAMESE_DOCX_FONT
    style._element.rPr.rFonts.set(qn('w:ascii'), VIETNAMESE_DOCX_FONT)
    style._element.rPr.rFonts.set(qn('w:hAnsi'), VIETNAMESE_DOCX_FONT)
    style._element.rPr.rFonts.set(qn('w:cs'), VIETNAMESE_DOCX_FONT)


def _add_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(fix_vietnamese_text(str(text or '')))
    run.font.name = VIETNAMESE_DOCX_FONT
    return p


def export_assessment_to_docx(assessment: Dict[str, Any], *, kind: str = '') -> Path:
    fd, out_path = tempfile.mkstemp(prefix='assessment_', suffix='.docx')
    os.close(fd)

    doc = Document()
    _set_times_new_roman(doc)
    _add_paragraph(doc, 'ĐỀ KIỂM TRA')

    title = str(assessment.get('title') or 'Assessment')
    level = str(assessment.get('level') or '')

    _add_paragraph(doc, f'Tiêu đề: {title}')
    if kind:
      _add_paragraph(doc, f'Loại: {kind}')
    if level:
      _add_paragraph(doc, f'Mức độ: {level}')
    _add_paragraph(doc, '')

    questions = assessment.get('questions') or []
    for idx, q in enumerate(questions, start=1):
        qtype = str(q.get('type') or '').lower()
        bloom = str(q.get('bloom_level') or '')
        stem = str(q.get('stem') or '')

        p = doc.add_paragraph()
        run = p.add_run(fix_vietnamese_text(f'Câu {idx} ({qtype.upper()})  Bloom: {bloom}\n'))
        run.bold = True
        run.font.name = VIETNAMESE_DOCX_FONT
        _add_paragraph(doc, stem)

        if qtype == 'mcq':
            opts = q.get('options') or []
            for oi, opt in enumerate(opts):
                label = chr(65 + oi)
                _add_paragraph(doc, f'{label}. {opt}', style='List Bullet')
        elif qtype == 'essay':
            mp = int(q.get('max_points') or 0)
            _add_paragraph(doc, f'(Tự luận) Điểm tối đa: {mp}')

        _add_paragraph(doc, '')

    doc.add_page_break()
    _add_paragraph(doc, 'ĐÁP ÁN / HƯỚNG DẪN CHẤM')

    for idx, q in enumerate(questions, start=1):
        qtype = str(q.get('type') or '').lower()
        if qtype == 'mcq':
            correct = q.get('correct_index')
            try:
                correct = int(correct)
                ans = chr(65 + correct)
            except Exception:
                ans = '?'
            _add_paragraph(doc, f'Câu {idx}: {ans}')
        elif qtype == 'essay':
            mp = int(q.get('max_points') or 0)
            _add_paragraph(doc, f'Câu {idx}: chấm theo rubric (tối đa {mp} điểm)')
            rubric = q.get('rubric') or []
            for r in rubric[:6]:
                try:
                    desc = str(r.get('criteria') or r.get('name') or '')
                    pts = r.get('points')
                    _add_paragraph(doc, f'- {desc}: {pts}')
                except Exception:
                    continue

    doc.save(out_path)
    return Path(out_path)


def export_multi_variant_docx(*, variants: List[Dict[str, Any]], title: str = 'Bộ đề') -> Path:
    fd, out_path = tempfile.mkstemp(prefix='assessment_multi_', suffix='.docx')
    os.close(fd)

    doc = Document()
    _set_times_new_roman(doc)
    _add_paragraph(doc, str(title or 'Bộ đề'))

    for idx, variant in enumerate(variants, start=1):
        code = str(variant.get('paper_code') or f'{idx:02d}')
        _add_paragraph(doc, f'Đề {code}')
        for q_idx, q in enumerate((variant.get('questions') or []), start=1):
            stem = str(q.get('stem') or '')
            qtype = str(q.get('type') or '').lower()
            _add_paragraph(doc, f'Câu {q_idx} ({qtype.upper()}): {stem}')
            if qtype == 'mcq':
                for oi, opt in enumerate((q.get('options') or []), start=0):
                    _add_paragraph(doc, f'{chr(65 + oi)}. {opt}', style='List Bullet')
        if idx < len(variants):
            doc.add_page_break()

    doc.add_page_break()
    _add_paragraph(doc, 'Đáp án tổng hợp')
    for idx, variant in enumerate(variants, start=1):
        code = str(variant.get('paper_code') or f'{idx:02d}')
        _add_paragraph(doc, f'Đề {code}')
        for q_idx, q in enumerate((variant.get('questions') or []), start=1):
            if str(q.get('type') or '').lower() == 'mcq':
                try:
                    answer = chr(65 + int(q.get('correct_index')))
                except Exception:
                    answer = '?'
                _add_paragraph(doc, f'Câu {q_idx}: {answer}')
            else:
                _add_paragraph(doc, f'Câu {q_idx}: Tự luận')

    doc.save(out_path)
    return Path(out_path)


def export_batch_to_zip(papers: List[Dict[str, Any]], include_answer_key: bool) -> Path:
    base_dir = Path(tempfile.mkdtemp(prefix='batch_exam_'))
    zip_path = base_dir / 'batch_exam.zip'

    metadata = {'total_papers': len(papers), 'codes': [str(p.get('paper_code') or '') for p in papers]}

    with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
        for idx, paper in enumerate(papers, start=1):
            code = str(paper.get('paper_code') or f'P{idx}')
            docx_path = export_assessment_to_docx(paper)
            zf.write(docx_path, arcname=f'paper_{code}.docx')

        if include_answer_key:
            key_doc = Document()
            _set_times_new_roman(key_doc)
            _add_paragraph(key_doc, 'BẢNG ĐÁP ÁN')
            for paper in papers:
                code = str(paper.get('paper_code') or '?')
                _add_paragraph(key_doc, f'Đề {code}')
                questions = paper.get('questions') or []
                for qi, q in enumerate(questions, start=1):
                    qtype = str(q.get('type') or '').lower()
                    if qtype == 'mcq':
                        try:
                            ans = chr(65 + int(q.get('correct_index')))
                        except Exception:
                            ans = '?'
                        _add_paragraph(key_doc, f'Q{qi}: {ans}')
                    else:
                        _add_paragraph(key_doc, f'Q{qi}: Essay')

            answer_key_path = base_dir / 'answer_key.docx'
            key_doc.save(answer_key_path)
            zf.write(answer_key_path, arcname='answer_key.docx')

        metadata_path = base_dir / 'metadata.json'
        metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding='utf-8')
        zf.write(metadata_path, arcname='metadata.json')

    return zip_path
