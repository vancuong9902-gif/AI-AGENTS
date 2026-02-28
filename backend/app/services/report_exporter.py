from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


def _register_vietnamese_font() -> str:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("DejaVu", path))
                return "DejaVu"
            except Exception:
                continue
    return "Helvetica"


def _build_level_chart(level_dist: dict[str, Any]) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(4, 4))
    labels = ["Giỏi", "Khá", "Trung bình", "Yếu"]
    keys = ["gioi", "kha", "trung_binh", "yeu"]
    values = [int(level_dist.get(k, 0) or 0) for k in keys]
    if sum(values) <= 0:
        values = [1, 0, 0, 0]

    ax.pie(values, labels=labels, colors=["#22c55e", "#3b82f6", "#f97316", "#ef4444"], autopct="%1.0f%%")
    ax.set_title("Phân bố xếp loại")
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def export_class_report_pdf(
    report_data: dict,
    output_path: str,
    *,
    class_name: str,
    teacher_name: str,
) -> str:
    font_name = _register_vietnamese_font()

    doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()
    vn_style = ParagraphStyle("vn", parent=styles["Normal"], fontName=font_name, fontSize=11, leading=15)
    title_style = ParagraphStyle(
        "title",
        parent=styles["Title"],
        fontName=font_name,
        textColor=colors.HexColor("#0f172a"),
        fontSize=18,
    )
    section_style = ParagraphStyle(
        "section",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=13,
        textColor=colors.HexColor("#111827"),
    )

    story: list[Any] = []
    story.append(Paragraph(f"BÁO CÁO KẾT QUẢ LỚP {class_name}", title_style))
    story.append(Paragraph(f"Giáo viên: {teacher_name}", vn_style))
    story.append(Spacer(1, 10))

    chart_buf = _build_level_chart(report_data.get("level_dist") or {})
    story.append(Image(chart_buf, width=210, height=210))
    story.append(Spacer(1, 10))

    story.append(Paragraph("NHẬN XÉT CỦA AI", section_style))
    story.append(Paragraph(str(report_data.get("narrative") or ""), vn_style))
    story.append(Spacer(1, 10))

    improvement = report_data.get("improvement") or {}
    story.append(
        Paragraph(
            f"Tiến bộ trung bình: {float(improvement.get('avg_delta') or 0.0):.1f} điểm • "
            f"Số học sinh tiến bộ: {int(improvement.get('improved_count') or 0)}",
            vn_style,
        )
    )
    story.append(Spacer(1, 8))

    weak_topics = report_data.get("weak_topics") or []
    if weak_topics:
        story.append(Paragraph("CHỦ ĐỀ CẦN CẢI THIỆN", section_style))
        for t in weak_topics[:5]:
            story.append(
                Paragraph(
                    f"• {t.get('topic', 'N/A')}: {float(t.get('avg_pct') or 0.0):.1f}% — {t.get('suggestion', '')}",
                    vn_style,
                )
            )
        story.append(Spacer(1, 10))

    student_data = [["Họ tên", "Đầu vào", "Cuối kỳ", "Tiến bộ", "Xếp loại"]]
    for s in report_data.get("students") or []:
        entry = float(s.get("entry_score") or 0.0)
        final = float(s.get("final_score") or 0.0)
        delta = final - entry
        student_data.append(
            [
                str(s.get("name") or "N/A"),
                f"{entry:.1f}",
                f"{final:.1f}",
                f"+{delta:.1f}" if delta >= 0 else f"{delta:.1f}",
                str(s.get("level") or "").upper(),
            ]
        )

    table = Table(student_data, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e5e7eb")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1d5db")),
                ("ALIGN", (1, 1), (-1, -1), "CENTER"),
            ]
        )
    )
    story.append(Paragraph("DANH SÁCH HỌC SINH", section_style))
    story.append(table)

    doc.build(story)
    return output_path


def export_class_report_docx(
    report_data: dict,
    output_path: str,
    *,
    class_name: str,
    teacher_name: str,
) -> str:
    doc = Document()
    doc.add_heading(f"BÁO CÁO KẾT QUẢ LỚP {class_name}", level=1)
    doc.add_paragraph(f"Giáo viên: {teacher_name}")

    chart_buf = _build_level_chart(report_data.get("level_dist") or {})
    doc.add_picture(chart_buf, width=Inches(3.2))

    doc.add_heading("NHẬN XÉT CỦA AI", level=2)
    doc.add_paragraph(str(report_data.get("narrative") or ""))

    improvement = report_data.get("improvement") or {}
    doc.add_paragraph(
        f"Tiến bộ trung bình: {float(improvement.get('avg_delta') or 0.0):.1f} điểm | "
        f"Số học sinh tiến bộ: {int(improvement.get('improved_count') or 0)}"
    )

    weak_topics = report_data.get("weak_topics") or []
    if weak_topics:
        doc.add_heading("CHỦ ĐỀ CẦN CẢI THIỆN", level=2)
        for t in weak_topics[:5]:
            doc.add_paragraph(
                f"- {t.get('topic', 'N/A')}: {float(t.get('avg_pct') or 0.0):.1f}% — {t.get('suggestion', '')}"
            )

    doc.add_heading("DANH SÁCH HỌC SINH", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Họ tên"
    hdr[1].text = "Đầu vào"
    hdr[2].text = "Cuối kỳ"
    hdr[3].text = "Tiến bộ"
    hdr[4].text = "Xếp loại"

    for s in report_data.get("students") or []:
        entry = float(s.get("entry_score") or 0.0)
        final = float(s.get("final_score") or 0.0)
        delta = final - entry
        row = table.add_row().cells
        row[0].text = str(s.get("name") or "N/A")
        row[1].text = f"{entry:.1f}"
        row[2].text = f"{final:.1f}"
        row[3].text = f"+{delta:.1f}" if delta >= 0 else f"{delta:.1f}"
        row[4].text = str(s.get("level") or "").upper()

    doc.save(output_path)
    return output_path


def make_export_path(*, classroom_id: int, extension: str) -> str:
    fd, out_path = tempfile.mkstemp(prefix=f"classroom_{int(classroom_id)}_report_", suffix=f".{extension}")
    os.close(fd)
    return str(Path(out_path))
